from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass, field
from datetime import date, datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Iterator


@dataclass
class Document:
    """One source document, before chunking."""
    doc_id: str
    text: str
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)


class _HTMLText(HTMLParser):
    SKIP = {"script", "style", "nav", "footer", "noscript"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in self.SKIP:
            self._skip_depth += 1
        elif tag in {"p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self.parts.append("\n")
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self.parts.append("#" * int(tag[1]) + " ")

    def handle_endtag(self, tag):
        if tag in self.SKIP and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data):
        if not self._skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        return re.sub(r"\n{3,}", "\n\n", "".join(self.parts)).strip()


def load_html(raw: str) -> str:
    p = _HTMLText()
    p.feed(raw)
    return p.text()


_FRONT_MATTER_RE = re.compile(r"\A---[ \t]*\r?\n(.*?)\r?\n---[ \t]*\r?\n", re.DOTALL)
_ISO_DATE_RE = re.compile(r"\A(\d{4})-(\d{2})-(\d{2})")


def split_front_matter(text: str) -> tuple[dict[str, str], str]:
    """Pull a leading `--- key: value ---` block off a document.

    Deliberately not YAML: the only thing this needs to read is a handful of
    flat scalars (`effective_date`, `supersedes`), and a real YAML parser is a
    dependency plus a parser that executes more of the file than intended.

    The block is *removed* from the returned text, not just read. Left in
    place, `_sections()` in chunking.py treats it as body content and produces
    a junk chunk with an empty heading -- which then competes for retrieval
    against real content.
    """
    match = _FRONT_MATTER_RE.match(text)
    if not match:
        return {}, text
    meta: dict[str, str] = {}
    for line in match.group(1).splitlines():
        key, sep, value = line.partition(":")
        if sep and key.strip():
            meta[key.strip()] = value.strip().strip("'\"")
    return meta, text[match.end():]


def parse_effective_date(value: str) -> str | None:
    """Accept an ISO date, reject anything else rather than guessing.

    A date that silently fails to parse would make a document look undated,
    which under recency weighting means "as old as possible" -- a document
    would sink in the ranking because of a typo in its header, with nothing
    reporting it.
    """
    m = _ISO_DATE_RE.match(value.strip())
    if not m:
        return None
    year, month, day = (int(g) for g in m.groups())
    try:
        date(year, month, day)
    except ValueError:
        return None
    return f"{year:04d}-{month:02d}-{day:02d}"


def load_csv(raw: str) -> str:
    """Render a CSV as one 'Field: value' record per block.

    Row-per-chunk beats naive text splitting here: a retrieved chunk stays a
    complete record instead of half of two rows.
    """
    reader = csv.DictReader(io.StringIO(raw))
    blocks = []
    for i, row in enumerate(reader):
        body = "\n".join(f"{k}: {v}" for k, v in row.items() if v not in (None, ""))
        blocks.append(f"### record {i + 1}\n{body}")
    return "\n\n".join(blocks)


def load_pdf(path: Path) -> str:
    """PDF text via pymupdf4llm, with automatic OCR for scanned pages.

    pymupdf4llm decides per page whether it has a usable text layer; pages that
    don't (scans, or a text layer with missing glyphs) fall back to OCR via
    RapidOCR -- a pure-Python engine (see requirements.txt) whose models ship
    inside the wheel, so this needs no system binary and no network call at
    ingest time, unlike a Tesseract-based setup. Clean, born-digital pages skip
    OCR entirely, so this stays fast on the common case. Output is Markdown
    with headings preserved, which is what structure-first chunking (the
    shipped default) is built to split on -- a scanned PDF chunks exactly as
    well as a hand-written .md file once OCR'd.

    If no OCR engine is importable at all (e.g. it was stripped from a minimal
    install), a scanned page silently yields no text -- fail loudly instead:
    a document that ingests as zero content is worse than one that doesn't
    ingest.
    """
    import pymupdf4llm
    text = pymupdf4llm.to_markdown(str(path))
    if not text.strip():
        raise RuntimeError(
            f"{path.name}: extracted no text at all. If this is a scanned PDF, "
            "check that `rapidocr` is installed (see requirements.txt) -- "
            "without an OCR engine, scanned pages silently produce nothing."
        )
    return text


def load_docx(path: Path) -> str:
    """Word text via python-docx, headings mapped to Markdown.

    Same reasoning as the PDF loader: structure-first chunking splits on
    Markdown headings, so a "Heading 2" paragraph style becomes "## text"
    rather than a plain line indistinguishable from body text. Tables render
    as tab-separated rows -- good enough to keep a row's cells together in
    one chunk without pulling in a Markdown-table renderer for what's usually
    a handful of small tables per document.
    """
    from docx import Document as _Docx

    doc = _Docx(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style == "title":
            parts.append(f"# {text}")
        elif style.startswith("heading"):
            level = next((c for c in style if c.isdigit()), "1")
            parts.append(f"{'#' * min(int(level), 6)} {text}")
        else:
            parts.append(text)
    for table in doc.tables:
        rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        rows = [r for r in rows if r.strip("\t")]
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def load_xlsx(path: Path) -> str:
    """Excel text, one record block per row per sheet -- the same shape as
    load_csv, extended with a sheet name since a workbook can have several."""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    blocks: list[str] = []
    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h).strip() if h is not None else "" for h in rows[0]]
        for i, row in enumerate(rows[1:], start=1):
            body = "\n".join(f"{h}: {v}" for h, v in zip(headers, row, strict=True)
                             if h and v not in (None, ""))
            if body:
                blocks.append(f"### {sheet.title} — record {i}\n{body}")
    return "\n\n".join(blocks)


LOADERS = {
    ".txt": lambda p: p.read_text(errors="replace"),
    ".md": lambda p: strip_html_comments(p.read_text(errors="replace")),
    ".markdown": lambda p: strip_html_comments(p.read_text(errors="replace")),
    ".html": lambda p: load_html(p.read_text(errors="replace")),
    ".htm": lambda p: load_html(p.read_text(errors="replace")),
    ".csv": lambda p: load_csv(p.read_text(errors="replace")),
    ".json": lambda p: json.dumps(json.loads(p.read_text()), indent=2),
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
}

SUPPORTED = tuple(LOADERS)

# Codepoints that render as nothing but survive copy-paste into a document:
# zero-width spaces/joiners, the bidi overrides, and the BOM. They are the
# standard way to hide an instruction from a human reviewer while leaving it
# perfectly legible to the model -- someone eyeballing the page in Word sees a
# clean paragraph, the retriever sees the payload. Written as hex ranges on
# purpose: a source file that itself contains invisible characters is neither
# reviewable nor greppable.
_INVISIBLE_TABLE = {
    cp: None
    for lo, hi in (
        (0x200B, 0x200F),   # zero-width space/non-joiner/joiner, LTR/RTL marks
        (0x202A, 0x202E),   # bidi embedding and override
        (0x2066, 0x2069),   # bidi isolates
        (0xFEFF, 0xFEFF),   # BOM / zero-width no-break space
    )
    for cp in range(lo, hi + 1)
}


# An HTML comment is invisible in rendered Markdown, so anything hidden in one is
# content no reader ever agreed to publish -- and a convenient place to park an
# instruction aimed at the model. The HTML loader already drops comments (HTMLParser
# ignores them unless handle_comment is implemented); Markdown read as raw text did
# not, which was an inconsistency, not a decision.
_HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)


def strip_html_comments(text: str) -> str:
    return _HTML_COMMENT_RE.sub("", text)


def strip_invisible(text: str) -> str:
    return text.translate(_INVISIBLE_TABLE)


def load_path(path: str | Path) -> Iterator[Document]:
    """Yield Documents from a file or (recursively) a directory."""
    p = Path(path)
    if not p.exists():
        # Fail loudly. A silent zero-document ingest is the worst possible
        # outcome: the index looks fine and every answer becomes a refusal.
        raise FileNotFoundError(f"nothing to ingest at {p}")
    files = [p] if p.is_file() else sorted(f for f in p.rglob("*") if f.is_file())
    if not files:
        raise FileNotFoundError(f"no files found under {p}")
    for f in files:
        ext = f.suffix.lower()
        if ext not in LOADERS:
            continue
        # Applied here rather than in each loader so every format -- including
        # ones added later -- gets it without anyone having to remember.
        text = strip_invisible(LOADERS[ext](f))
        front, text = split_front_matter(text)
        if not text.strip():
            continue

        meta: dict[str, Any] = {"filename": f.name, "ext": ext, "bytes": f.stat().st_size}
        raw_date = front.get("effective_date", "")
        effective = parse_effective_date(raw_date) if raw_date else None
        if raw_date and effective is None:
            raise ValueError(
                f"{f.name}: effective_date {raw_date!r} is not an ISO date (YYYY-MM-DD). "
                "Leaving it unparsed would make the document look undated, which under "
                "recency weighting means 'as old as possible'."
            )
        if effective:
            meta["effective_date"] = effective
            meta["date_source"] = "front_matter"
        else:
            # mtime is a weak fallback and the README says so: git does not
            # preserve it, so a fresh clone stamps every file with checkout
            # time. Good enough for a live client folder, useless for a repo.
            meta["effective_date"] = datetime.fromtimestamp(
                f.stat().st_mtime, tz=timezone.utc).date().isoformat()
            meta["date_source"] = "mtime"
        if front.get("supersedes"):
            meta["supersedes"] = front["supersedes"]

        yield Document(doc_id=f.as_posix(), text=text, source=f.as_posix(), metadata=meta)
