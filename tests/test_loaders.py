import pymupdf
import pytest
from docx import Document
from openpyxl import Workbook

from app.ingest.loaders import load_docx, load_pdf, load_xlsx


def _text_pdf(path, heading: str, body: str) -> None:
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), heading, fontsize=18)
    page.insert_text((72, 110), body, fontsize=11)
    doc.save(str(path))
    doc.close()


def _scanned_pdf(path, heading: str, body: str) -> None:
    """A PDF with no text layer at all -- an image of text, like a real scan."""
    src = pymupdf.open()
    src_page = src.new_page()
    src_page.insert_text((72, 72), heading, fontsize=18)
    src_page.insert_text((72, 110), body, fontsize=11)
    pixmap = src_page.get_pixmap(dpi=200)
    image_bytes = pixmap.tobytes("png")
    src.close()

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_image(page.rect, stream=image_bytes)
    doc.save(str(path))
    doc.close()


def test_load_pdf_extracts_text_and_keeps_the_heading(tmp_path):
    pdf_path = tmp_path / "policy.pdf"
    _text_pdf(pdf_path, "Refund Policy", "Refunds are issued within 30 days of purchase.")

    text = load_pdf(pdf_path)

    assert "Refund Policy" in text
    assert "30 days" in text
    assert text.strip().startswith("#")  # structure-first chunking splits on this


def test_load_pdf_ocrs_a_scanned_page_with_no_text_layer(tmp_path):
    pdf_path = tmp_path / "scan.pdf"
    _scanned_pdf(pdf_path, "International Shipping", "We ship to 40 countries.")

    text = load_pdf(pdf_path)

    assert "International Shipping" in text
    assert "40 countries" in text


def test_load_pdf_fails_loudly_on_a_genuinely_blank_page(tmp_path):
    pdf_path = tmp_path / "blank.pdf"
    doc = pymupdf.open()
    doc.new_page()
    doc.save(str(pdf_path))
    doc.close()

    with pytest.raises(RuntimeError, match="extracted no text"):
        load_pdf(pdf_path)


def test_load_docx_maps_heading_styles_to_markdown(tmp_path):
    docx_path = tmp_path / "warranty.docx"
    doc = Document()
    doc.add_heading("Warranty Policy", level=1)
    doc.add_paragraph("Intro paragraph about warranty terms.")
    doc.add_heading("Batteries", level=2)
    doc.add_paragraph("Batteries are covered for 12 months.")
    doc.save(str(docx_path))

    text = load_docx(docx_path)

    assert "# Warranty Policy" in text
    assert "## Batteries" in text
    assert "12 months" in text


def test_load_docx_keeps_table_rows_together(tmp_path):
    docx_path = tmp_path / "coverage.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Coverage"
    table.cell(1, 0).text = "Battery"
    table.cell(1, 1).text = "12 months"
    doc.save(str(docx_path))

    text = load_docx(docx_path)

    assert "Battery" in text and "12 months" in text
    # same row, not split across two chunks worth of text
    battery_line = next(line for line in text.splitlines() if "Battery" in line)
    assert "12 months" in battery_line


def test_load_xlsx_emits_one_record_block_per_row(tmp_path):
    xlsx_path = tmp_path / "warranty.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Warranty"
    ws.append(["Item", "Coverage"])
    ws.append(["Battery", "12 months"])
    ws.append(["Screen", "24 months"])
    wb.save(str(xlsx_path))

    text = load_xlsx(xlsx_path)

    assert "Warranty" in text
    assert "Item: Battery" in text
    assert "Coverage: 12 months" in text
    assert "Item: Screen" in text
    assert text.count("record") == 2


def test_load_xlsx_covers_every_sheet(tmp_path):
    xlsx_path = tmp_path / "multi.xlsx"
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Warranty"
    ws1.append(["Item"])
    ws1.append(["Battery"])
    ws2 = wb.create_sheet("Shipping")
    ws2.append(["Region"])
    ws2.append(["EU"])
    wb.save(str(xlsx_path))

    text = load_xlsx(xlsx_path)

    assert "Warranty" in text and "Battery" in text
    assert "Shipping" in text and "EU" in text
